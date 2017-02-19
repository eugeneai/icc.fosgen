from docx import Document
from docx.shared import Inches


class Generator(object):

    def __init__(self, comps, table9, template=None):
        self.comps = comps
        self.table9 = table9
        self.doc = Document(template)

    def header(self, text, level):
        self.doc.add_heading(text, level)

    def add_results(self, header=False):
        if header:
            self.header("Результаты освоения дисциплины", 2)
        vals = list(self.table9.items())
        vals.sort()
        p = self.doc.add_paragraph(
            "В результате изучения дисциплины студенты должны:")
        lvals = len(vals)
        for sect in ["знать", "уметь", "владеть"]:
            p = self.doc.add_paragraph("")
            r = p.add_run(sect + ":")
            r.bold = True
            r.italic = True
            for row, zun in enumerate(vals):
                s = zun[1][0][1].strip()
                # ListNumber IntenseQuote
                if s.startswith(sect):
                    s = s.replace(sect, "", 1)
                    if row == lvals - 1:
                        end = '.'
                    else:
                        end = ";"
                    self.doc.add_paragraph(s + end, style='ListBullet')

    def add_table9_WP(self, header=True, page_breaks=True):
        # self.doc.add_paragraph(
        #    "(Таблица размещается на отдельном листе в Landscape)")
        if page_breaks:
            self.doc.add_page_break()
        if header:
            self.doc.add_paragraph(
                "Таблица 9 – Контролируемые элементы "
                "содержания дисциплины и виды учебных работ, "
                "по результатам выполнения которых и отчета "
                "по ним осуществляется текущий контроль")
        t9 = self.doc.add_table(rows=len(self.table9) + 3, cols=4 + 3 * 4)
        t9.style = 'TableGrid'
        cs = t9.rows[0].cells
        ls = t9.rows[2].cells
        cs[0].merge(ls[0])
        cs[0].text = "№\nп/п"
        cs[1].merge(ls[1])
        cs[1].text = "Контролируемые элементы содержания дисциплины"
        cs[2].merge(ls[2])
        cs[2].text = "Компетенции"
        cs[3].merge(ls[3])
        cs[3].text = "№ раздела, темы\nпо табл. 2"

        cs[4].merge(cs[-1])
        cs[4].text = "Текущий контроль успеваемости (ТК)"

        off = 4
        tk = t9.rows[2].cells
        tk2 = t9.rows[1].cells
        for i, iname in enumerate(["ЛР № по\nтабл. 3", "ПЗ/СЕМ №\nпо табл.4", "СРС № по\nтабл.5", "КП (КР) №\nпо табл.7"]):
            a = off + i * 3
            b = a + 2
            tk2[a].merge(tk2[b])
            tk2[a].text = iname
            for j in range(3):
                tk[off + i * 3 + j].text = "ТК № {}".format(j + 1)

        t9rows = list(self.table9.items())
        t9rows.sort()
        for k, t9row in t9rows:
            f5, rest = t9row[0], t9row[1:]

            comp = f5[2].strip().replace(" ", "\n")
            f5[2] = comp

            off = 0
            c = t9.rows[k + 2].cells
            for i, it in enumerate(f5):
                c[off + i].text = it
            off = 4
            x = off
            for i, ig in enumerate(rest):
                for j, jt in enumerate(ig):
                    c[x].text = jt
                    x += 1

        if page_breaks:
            self.doc.add_page_break()

    def add_table1_FOS(self, header=True, page_breaks=True, semesters=["XXX"]):
        if page_breaks:
            self.doc.add_page_break()

        t1 = self.doc.add_table(rows=len(self.comps) + 1, cols=5)
        t1.style = 'TableGrid'

        h = t1.rows[0].cells

        headers = """№\nп/п|
        Контролируемые\nкомпетенции|
        Контролируемые элементы содержания дисциплины|
        КИМ|
        Вид ФОС\n(текущий контроль №1,\nтекущий контроль №2,\nтекущий контроль №3,\nпромежуточная аттестация)"""

        headers = headers.split("|\n")

        [setattr(c, "text", t.strip()) for c, t in zip(h, headers)]

        off = 1
        for i, it in enumerate(self.comps.items()):
            code, name = it
            c = t1.rows[i + off].cells
            c[0].text = str(i + 1)
            c[1].text = "{} ({})".format(name, code)
            for sem in semesters:
                p3 = c[3].paragraphs[0]
                p3.add_run("Защита лабораторных работ,\n"
                           "Контрольные работы и вопросы,\n"
                           "Защита курсового проекта\n")
                p3.add_run("\n")
                p4 = c[4].paragraphs[0]
                p4.add_run("{} семестр\n".format(sem)).bold = True
                for tk in range(3):
                    p4.add_run("Текущий контроль № {}\n".format(tk + 1))
                p4.add_run("Промежуточная аттестация\n\n")

            p2 = c[2].paragraphs[0]
            for r9 in self.table9.values():
                f = r9[0]
                _, zname, opk = f[:3]
                pks = [o.strip() for o in opk.split(" ")]
                if code in pks:
                    p2.add_run(zname + ";\n")

        if page_breaks:
            self.doc.add_page_break()

    def add_table2_FOS(self,
                       header=True,
                       variant="на весь срок изучения дисциплины",
                       page_breaks=False,
                       input=True,
                       intermediate=True,
                       test="Экзамен",
                       weeks=None,
                       scores=None,
                       semesters=["XXX"]
                       ):
        if page_breaks:
            self.doc.add_page_break()

        if header:
            self.doc.add_paragraph(
                "Таблица 2 – План проведения оценочных мероприятий " + variant)

        items = []
        items.append(["Вид ФОС",
                      "Исходные требования к уровню усвоения",
                      "Объект\nоценивания",
                      "Вид контроля\n(все виды контроля, "
                      "используемые в ходе освоения дисциплины)",
                      "Период\nоценивания",
                      "Распределение общего кол-ва баллов"
                      ])
        for sem in semesters:
            items.append("{} семестр".format(sem))
            if input:
                items.append(["Входной\nконтроль",
                              "Рабочая\nпрограмма\nдисциплины",
                              "Конкретизированные результаты"
                              "предшествующего обучения",
                              "Опрос",
                              ])
            jpoint = len(items)
            for i in range(3):
                items.append(["Текущий контроль № {}".format(i + 1),
                              "Рабочая\nпрограмма\nдисциплины",
                              "Конкретизированные результаты обучения",
                              "Защита отчетов по лабораторным работам"
                              ])

            if intermediate:
                items.append(["Промежуточная аттестация",
                              "Рабочая\nпрограмма\nдисциплины",
                              "Обобщенные результаты обучения по дисциплине",
                              test,
                              "Согласно календарному графику учебного процесса"
                              ])

        t2 = self.doc.add_table(rows=len(items) + 1, cols=6)
        t2.style = 'TableGrid'

        t2.rows[jpoint].cells[1].merge(t2.rows[jpoint + 2].cells[1])
        t2.rows[jpoint].cells[2].merge(t2.rows[jpoint + 2].cells[2])

        off = 0
        for rn, item in enumerate(items):
            c = t2.rows[rn].cells
            if isinstance(item, str):
                c[0].merge(c[-1])
                c[0].paragraphs[0].add_run(item).bold = True
                continue
            for j, jt in enumerate(item):
                c[j].text = jt
            if scores and rn > 1:
                try:
                    srn = scores[rn - 2]
                    c[-1].text = "" if srn is None else "{}%".format(srn)
                except IndexError:
                    pass

            if weeks and rn > 1:
                try:
                    srn = weeks[rn - 2]
                    c[-2].text = "" if srn is None else "{}-я неделя".format(
                        srn)
                except IndexError:
                    pass

        r = t2.rows[-1].cells
        r[0].merge(r[-2])
        r[0].text = "ИТОГО"
        r[-1].text = "100%"
        if page_breaks:
            self.doc.add_page_break()

    def save(self, filename):
        return self.doc.save(filename)
